from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
ICON = ROOT / "assets" / "app_icon_source.png"
TEAL = "168B8C"
LIGHT_TEAL = "DDF3F2"
DARK = RGBColor(35, 55, 65)


def set_font(run, name: str = "Yu Gothic", size: float | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def style_document(path: Path, label: str) -> None:
    doc = Document(path)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("Title", 20, RGBColor(20, 110, 112)),
        ("Heading 1", 17, RGBColor(20, 110, 112)),
        ("Heading 2", 13, RGBColor(42, 105, 125)),
        ("Heading 3", 11.5, RGBColor(65, 85, 95)),
    ):
        style = styles[name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 4)
        style.paragraph_format.space_after = Pt(6)

    first = doc.paragraphs[0]
    icon_paragraph = first.insert_paragraph_before()
    icon_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    icon_paragraph.add_run().add_picture(str(ICON), width=Inches(1.0))
    first.style = styles["Title"]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in first.runs:
        set_font(run, size=20)

    subtitle = first.insert_paragraph_before()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"バージョン 0.2.0  |  {label}")
    set_font(run, size=9.5)
    run.font.color.rgb = RGBColor(90, 110, 120)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_font(run)
        if paragraph.style.name == "Source Code":
            paragraph.paragraph_format.left_indent = Cm(0.4)
            paragraph.paragraph_format.right_indent = Cm(0.4)
            for run in paragraph.runs:
                set_font(run, "Menlo", 9)

    for table in doc.tables:
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            if row_index == 0:
                repeat_header_row(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if row_index == 0:
                    shade_cell(cell, TEAL)
                elif row_index % 2 == 0:
                    shade_cell(cell, LIGHT_TEAL)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        set_font(run, size=9.5)
                        if row_index == 0:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_paragraph.add_run(f"Codex App かんたん切り替え  |  {label}")
    set_font(footer_run, size=8.5)
    footer_run.font.color.rgb = RGBColor(105, 120, 128)

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.save(path)


def build(markdown_name: str, output_name: str, label: str) -> None:
    source = DOCS / markdown_name
    output = DOCS / output_name
    subprocess.run(
        [
            "pandoc",
            str(source),
            "--from",
            "gfm",
            "--to",
            "docx",
            "--output",
            str(output),
        ],
        check=True,
    )
    style_document(output, label)


def main() -> None:
    build("SPECIFICATION.md", "Codex_App_かんたん切り替え_仕様書.docx", "仕様書")
    build("USER_MANUAL.md", "Codex_App_かんたん切り替え_操作マニュアル.docx", "操作マニュアル")


if __name__ == "__main__":
    main()
